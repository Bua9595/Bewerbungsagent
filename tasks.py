#!/usr/bin/env python3
"""
Einfache Task-CLI für den Bewerbungsagenten.

Beispiele:
  python tasks.py env-check
  python tasks.py gen-templates
  python tasks.py start
  python tasks.py open
  python tasks.py email-test
  python tasks.py list
  python tasks.py mail-list
  python tasks.py mail-list --dry-run
  python tasks.py tracker-sync
  python tasks.py mark-applied <job_uid>
  python tasks.py mark-ignored --url <link>
  python tasks.py prepare-applications --force-all
"""

import argparse
import os
import re
import json
import shutil
import subprocess
import sys
from datetime import datetime, timezone
from pathlib import Path
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication


def cmd_env_check(_args=None):
    from config import config
    print("=== ENV/Cfg ===")
    print("Sender:", config.SENDER_EMAIL or "<leer>")
    print("SMTP:", config.SMTP_SERVER or "<leer>", config.SMTP_PORT)
    print("Recipients:", config.RECIPIENT_EMAILS or [])
    print(
        "Profile:",
        getattr(config, "PROFILE_NAME", ""),
        getattr(config, "PROFILE_EMAIL", ""),
    )
    pwd_len = len(config.SENDER_PASSWORD or "")
    print(
        "Password set:",
        "Yes" if pwd_len else "No",
        "(len=",
        pwd_len,
        ")",
    )


def cmd_gen_templates(_args=None):
    from direkt_job_finder import DirectJobFinder
    app = DirectJobFinder()
    app.save_application_templates()
    app.create_job_tracking_sheet()
    print("Templates/Tracking aktualisiert.")


def cmd_start(_args=None):
    from direkt_job_finder import DirectJobFinder
    DirectJobFinder().run_complete_job_hunt()


def cmd_open(_args=None):
    from direkt_job_finder import DirectJobFinder
    DirectJobFinder().open_job_portals_automatically()


def cmd_email_test(_args=None):
    # nutzt das getestete Testskript, das (bool, lines) zurückgibt
    from test_email_config import test_email_connection

    success, output_lines = test_email_connection()
    for line in output_lines:
        print(line)

    raise SystemExit(0 if success else 1)


def cmd_list(_args=None):
    from job_collector import collect_jobs, export_csv, export_json
    jobs = collect_jobs()
    if not jobs:
        print("Keine Treffer. CSV/Mail übersprungen.")
        return
    export_csv(jobs)
    export_json(jobs)
    for i, j in enumerate(jobs[:20], 1):
        company = j.company
        location = j.location
        if (not company or not location) and (j.raw_title or j.title):
            from job_collector import _extract_from_multiline_title

            t2, c2, l2 = _extract_from_multiline_title(j.raw_title or j.title)
            if t2:
                j.title = t2
            if not company and c2:
                company = c2
            if not location and l2:
                location = l2

        company = company or "Firma unbekannt"
        location = location or "Ort unbekannt"
        print(f"{i:02d}. [{j.match:^5}] {j.title} - {company} - {location}")
        print(f"    {j.link}")


def cmd_mail_list(args=None):
    """
    Sendet Job-Alert per Mail (und ggf. WhatsApp).
    Neue Jobs plus Erinnerungen fuer offene Jobs.
    """
    try:
        from job_collector import collect_jobs, export_json
        from email_automation import email_automation
        from job_state import (
            OPEN_STATUSES,
            TERMINAL_STATUSES,
            STATUS_APPLIED,
            STATUS_CLOSED,
            STATUS_IGNORED,
            STATUS_NEW,
            STATUS_NOTIFIED,
            build_job_uid,
            canonicalize_url,
            load_state,
            now_iso,
            parse_ts,
            save_state,
            should_send_reminder,
        )
        from job_tracker import (
            apply_tracker_marks,
            get_tracker_path,
            load_tracker,
            write_tracker,
        )
        from logger import job_logger
    except Exception as e:
        print(f"Mail-Liste Fehler: {e}")
        return

    stamp = now_iso()
    now_dt = parse_ts(stamp) or datetime.now(timezone.utc)

    state_path = Path("generated/job_state.json")
    seen_path = Path("generated/seen_jobs.json")
    migrated_from_seen = (not state_path.exists()) and seen_path.exists()

    state = load_state(now=stamp)

    tracker_path = get_tracker_path()
    tracker_rows = load_tracker(tracker_path)
    tracker_updates = apply_tracker_marks(state, tracker_rows)

    rows = collect_jobs()
    scraped_total = len(rows)
    if not rows:
        applied_count = sum(
            1 for record in state.values() if record.get("status") == STATUS_APPLIED
        )
        ignored_count = sum(
            1 for record in state.values() if record.get("status") == STATUS_IGNORED
        )
        stats = {
            "scraped_total": scraped_total,
            "unique_total": 0,
            "state_total": len(state),
            "newly_added": 0,
            "active_seen_this_run": 0,
            "mailed_new_count": 0,
            "mailed_reminder_count": 0,
            "marked_closed_count": 0,
            "applied_count": applied_count,
            "ignored_count": ignored_count,
            "dry_run": bool(args and getattr(args, "dry_run", False)),
            "mail_sent": False,
        }

        print("Mail-Statistik:")
        for key in [
            "scraped_total",
            "unique_total",
            "state_total",
            "newly_added",
            "active_seen_this_run",
            "mailed_new_count",
            "mailed_reminder_count",
            "marked_closed_count",
            "applied_count",
            "ignored_count",
        ]:
            print(f"{key}: {stats[key]}")

        job_logger.info(
            "Mail-Statistik "
            + ", ".join(f"{k}={v}" for k, v in stats.items())
        )
        if migrated_from_seen or tracker_updates:
            save_state(state)
        if migrated_from_seen:
            print("Hinweis: seen_jobs.json wurde in job_state.json migriert.")
        write_tracker(state, tracker_path, tracker_rows)
        return
    export_json(rows)

    min_score = int(os.getenv("MIN_SCORE_MAIL", "2") or 2)

    # Nur Score-Filter (match kann durch Heuristik/Normalisierung mal leer sein)
    filtered = [r for r in rows if (r.score or 0) >= min_score]

    # Fallback: wenn Filter leer, nimm Top 10 statt silent skip
    if not filtered:
        filtered = rows[:10]

    payload = [
        r.__dict__ if hasattr(r, "__dict__") else dict(r)
        for r in filtered
    ]
    unique_total = len(payload)


    reminder_days = int(os.getenv("REMINDER_DAYS", "2") or 2)
    close_missing_runs = int(os.getenv("CLOSE_MISSING_RUNS", "3") or 3)
    close_not_seen_days = int(os.getenv("CLOSE_NOT_SEEN_DAYS", "7") or 7)
    daily_reminders = str(os.getenv("REMINDER_DAILY", "false")).lower() in {
        "1",
        "true",
        "t",
        "yes",
        "y",
        "ja",
        "j",
    }

    seen_this_run = set()
    newly_added = 0

    def _score_val(value):
        try:
            return float(value)
        except Exception:
            return 0

    for r in payload:
        job_uid, canonical_url = build_job_uid(r)
        seen_this_run.add(job_uid)

        link = (
            r.get("link")
            or r.get("url")
            or r.get("apply_url")
            or r.get("applyLink")
            or ""
        )

        record = state.get(job_uid)
        if not record:
            record = {
                "job_uid": job_uid,
                "source": r.get("source") or "",
                "canonical_url": canonical_url or canonicalize_url(link) or link,
                "link": link,
                "title": r.get("title") or "",
                "company": r.get("company") or "",
                "location": r.get("location") or "",
                "first_seen_at": stamp,
                "last_seen_at": stamp,
                "last_sent_at": None,
                "status": STATUS_NEW,
                "score": r.get("score", ""),
                "match": r.get("match", ""),
                "date": r.get("date", ""),
                "missing_runs": 0,
            }
            state[job_uid] = record
            newly_added += 1
            continue

        record["source"] = r.get("source") or record.get("source", "")
        record["canonical_url"] = (
            canonical_url
            or record.get("canonical_url", "")
            or canonicalize_url(link)
            or link
        )
        record["link"] = link or record.get("link", "")
        record["title"] = r.get("title") or record.get("title", "")
        record["company"] = r.get("company") or record.get("company", "")
        record["location"] = r.get("location") or record.get("location", "")
        if r.get("score") not in (None, ""):
            record["score"] = r.get("score")
        if r.get("match"):
            record["match"] = r.get("match")
        if r.get("date"):
            record["date"] = r.get("date")
        record["last_seen_at"] = stamp
        record["missing_runs"] = 0

        status = record.get("status") or STATUS_NEW
        if status == STATUS_CLOSED:
            status = STATUS_NOTIFIED if record.get("last_sent_at") else STATUS_NEW
        if status not in (STATUS_APPLIED, STATUS_IGNORED):
            record["status"] = status

    marked_closed_count = 0
    for uid, record in state.items():
        if uid in seen_this_run:
            continue
        if record.get("status") in (STATUS_APPLIED, STATUS_IGNORED, STATUS_CLOSED):
            continue
        record["missing_runs"] = int(record.get("missing_runs", 0)) + 1
        last_seen = parse_ts(record.get("last_seen_at"))
        days_missing = (now_dt - last_seen).days if last_seen else 0
        if (
            close_missing_runs > 0
            and record["missing_runs"] >= close_missing_runs
        ) or (
            close_not_seen_days > 0
            and days_missing >= close_not_seen_days
        ):
            record["status"] = STATUS_CLOSED
            marked_closed_count += 1

    new_jobs = []
    reminder_jobs = []
    new_uids = set()

    for uid in seen_this_run:
        record = state.get(uid)
        if not record:
            continue
        if record.get("status") in TERMINAL_STATUSES:
            continue
        if record.get("status") == STATUS_NEW:
            new_jobs.append(record)
            new_uids.add(uid)

    for uid in seen_this_run:
        record = state.get(uid)
        if not record:
            continue
        if record.get("status") in TERMINAL_STATUSES:
            continue
        if (
            record.get("status") in OPEN_STATUSES
            and should_send_reminder(
                record.get("last_sent_at"),
                now_dt,
                reminder_days,
                daily_reminders,
            )
        ):
            if uid not in new_uids:
                reminder_jobs.append(record)

    new_jobs.sort(key=lambda r: _score_val(r.get("score")), reverse=True)
    reminder_jobs.sort(key=lambda r: _score_val(r.get("score")), reverse=True)

    active_seen_this_run = sum(
        1
        for uid in seen_this_run
        if uid in state and state[uid].get("status") not in TERMINAL_STATUSES
    )
    applied_count = sum(
        1 for record in state.values() if record.get("status") == STATUS_APPLIED
    )
    ignored_count = sum(
        1 for record in state.values() if record.get("status") == STATUS_IGNORED
    )

    mailed_new_count = 0
    mailed_reminder_count = 0
    mail_sent = False

    if new_jobs or reminder_jobs:
        if args and getattr(args, "dry_run", False):
            mailed_new_count = len(new_jobs)
            mailed_reminder_count = len(reminder_jobs)
            print(
                f"[DRY RUN] Mail waere gesendet worden "
                f"({mailed_new_count} neu, {mailed_reminder_count} Reminder)."
            )
        else:
            ok = email_automation.send_job_alert(new_jobs, reminder_jobs)
            if ok:
                mail_sent = True
                mailed_new_count = len(new_jobs)
                mailed_reminder_count = len(reminder_jobs)
                for record in new_jobs + reminder_jobs:
                    record["status"] = STATUS_NOTIFIED
                    record["last_sent_at"] = stamp
                print(
                    f"E-Mail gesendet ({mailed_new_count} neu, "
                    f"{mailed_reminder_count} Reminder)"
                )
            else:
                print("Mail/WhatsApp uebersprungen (disabled oder Fehler).")
    else:
        print("Keine neuen oder offenen Jobs zum Senden.")

    save_state(state)

    stats = {
        "scraped_total": scraped_total,
        "unique_total": unique_total,
        "state_total": len(state),
        "newly_added": newly_added,
        "active_seen_this_run": active_seen_this_run,
        "mailed_new_count": mailed_new_count,
        "mailed_reminder_count": mailed_reminder_count,
        "marked_closed_count": marked_closed_count,
        "applied_count": applied_count,
        "ignored_count": ignored_count,
        "dry_run": bool(args and getattr(args, "dry_run", False)),
        "mail_sent": mail_sent,
    }

    print("Mail-Statistik:")
    for key in [
        "scraped_total",
        "unique_total",
        "state_total",
        "newly_added",
        "active_seen_this_run",
        "mailed_new_count",
        "mailed_reminder_count",
        "marked_closed_count",
        "applied_count",
        "ignored_count",
    ]:
        print(f"{key}: {stats[key]}")

    job_logger.info(
        "Mail-Statistik "
        + ", ".join(f"{k}={v}" for k, v in stats.items())
    )
    if migrated_from_seen:
        print("Hinweis: seen_jobs.json wurde in job_state.json migriert.")
    write_tracker(state, tracker_path, tracker_rows)

def cmd_tracker_sync(_args=None):
    from job_state import load_state, save_state
    from job_tracker import (
        apply_tracker_marks,
        get_tracker_path,
        load_tracker,
        write_tracker,
    )

    state = load_state()
    if not state:
        print("Kein job_state.json vorhanden.")
        return

    tracker_path = get_tracker_path()
    tracker_rows = load_tracker(tracker_path)
    if not tracker_rows:
        print("Kein job_tracker vorhanden.")
        return

    updates = apply_tracker_marks(state, tracker_rows)
    if updates:
        save_state(state)
    write_tracker(state, tracker_path, tracker_rows)
    print(f"Tracker Sync: {updates} Aktualisierungen.")


def _resolve_job_uid(state, job_uid, url):
    from job_state import canonicalize_url

    if job_uid:
        matches = [uid for uid in state.keys() if uid.startswith(job_uid)]
        if len(matches) == 1:
            return matches[0]
        if len(matches) > 1:
            print("Mehrere Treffer fuer job_uid, bitte laenger angeben.")
            return None
        print("Keine Treffer fuer job_uid.")
        return None

    if url:
        target = url.strip()
        target_canon = canonicalize_url(target)
        matches = []
        for uid, record in state.items():
            link = record.get("link") or ""
            canon = record.get("canonical_url") or ""
            if target and target == link:
                matches.append(uid)
                continue
            if target_canon:
                if canon and canonicalize_url(canon) == target_canon:
                    matches.append(uid)
                    continue
                if link and canonicalize_url(link) == target_canon:
                    matches.append(uid)
        if len(matches) == 1:
            return matches[0]
        if len(matches) > 1:
            print("Mehrere Treffer fuer URL, bitte job_uid nutzen.")
            return None
        print("Keine Treffer fuer URL.")
        return None

    print("Bitte job_uid oder --url angeben.")
    return None


def _mark_job_state(args, status):
    from job_state import load_state, save_state

    state = load_state()
    if not state:
        print("Kein job_state.json vorhanden.")
        return

    job_uid = _resolve_job_uid(state, getattr(args, "job_uid", ""), getattr(args, "url", ""))
    if not job_uid:
        return

    record = state.get(job_uid)
    if not record:
        print("Job nicht gefunden.")
        return

    prev = record.get("status") or ""
    record["status"] = status
    save_state(state)
    try:
        from job_tracker import get_tracker_path, load_tracker, write_tracker

        tracker_path = get_tracker_path()
        tracker_rows = load_tracker(tracker_path)
        write_tracker(state, tracker_path, tracker_rows)
    except Exception:
        pass

    title = record.get("title") or "Titel unbekannt"
    company = record.get("company") or "Firma unbekannt"
    print(f"{job_uid}: {prev} -> {status} ({title} - {company})")


def cmd_mark_applied(args):
    from job_state import STATUS_APPLIED

    _mark_job_state(args, STATUS_APPLIED)


def cmd_mark_ignored(args):
    from job_state import STATUS_IGNORED

    _mark_job_state(args, STATUS_IGNORED)


def cmd_archive_sent(args):
    """
    Kopiert ein versendetes Anschreiben nach 04_Versendete_Bewerbungen/<Firma>/ (oder --copy-sent-dir).
    Erwartet eine bestehende DOCX (z.B. aus out/), optional Firmenname override.
    """
    src = Path(args.file).expanduser()
    if not src.exists():
        print(f"FEHLER: Datei nicht gefunden: {src}")
        raise SystemExit(1)

    company = args.company
    if not company:
        stem = src.stem
        parts = stem.split("_")
        if parts:
            company = parts[0]
    company = company or "Unbekannt"

    dest_base = Path(args.dest).expanduser() if args.dest else Path.cwd() / "04_Versendete_Bewerbungen"
    dest_dir = dest_base / _sanitize_filename(company)
    dest_dir.mkdir(parents=True, exist_ok=True)
    dest = dest_dir / src.name
    shutil.copy2(src, dest)
    print(f"Kopie erstellt: {dest}")


def cmd_verify(_args=None):
    """
    Lightweight Verify: config check, compileall, presence of key dirs/files.
    """
    ok = True
    try:
        from config import config
        config.validate_config()
        print("Config: OK")
    except Exception as e:
        ok = False
        print(f"Config-Check fehlgeschlagen: {e}")

    required_dirs = ["Anschreiben_Templates", "out", "data"]
    for d in required_dirs:
        if not Path(d).exists():
            ok = False
            print(f"FEHLT: {d}")
    for tpl in ["T1_ITSup.docx", "T2_Systemtechnik.docx", "T3_Logistik.docx"]:
        if not (Path("Anschreiben_Templates") / tpl).exists():
            ok = False
            print(f"Template fehlt: Anschreiben_Templates/{tpl}")

    try:
        subprocess.check_call(
            [sys.executable, "-m", "compileall", "."],
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )
        print("compileall: OK")
    except Exception as e:
        ok = False
        print(f"compileall fehlgeschlagen: {e}")

    if ok:
        print("Verify: OK")
    else:
        print("Verify: FEHLER")


# ---------------------------
# prepare-applications LOGIK
# ---------------------------

_COMPANY_HINT_RE = re.compile(
    r"\b(ag|gmbh|sa|s\.a\.|kg|sarl|s\u00e0rl|sarl\.?|ltd|inc|llc)\b",
    re.IGNORECASE,
)

_LABEL_RE = re.compile(
    r"(arbeitsort|pensum|vertragsart|einfach bewerben|neu)",
    re.IGNORECASE,
)

_RELDATE_INLINE_RE = re.compile(
    r"\b(heute|gestern|vorgestern|letzte woche|letzten monat|vor \d+ (stunden?|tagen|wochen|monaten?))\b",
    re.IGNORECASE,
)

_CITY_HINT_RE = re.compile(
    r"\b("
    r"z\u00fcrich|zurich|zuerich|"
    r"b\u00fclach|buelach|"
    r"kloten|winterthur|baden|zug|aarau|basel|bern|luzern|thun|"
    r"gen\u00e8ve|geneve|"
    r"schweiz"
    r")\b",
    re.IGNORECASE,
)


def _sanitize_filename(s: str) -> str:
    s = s.strip()
    s = re.sub(r"[\\/:*?\"<>|]", "_", s)
    s = re.sub(r"\s+", " ", s)
    return s[:120] if len(s) > 120 else s


def _normalize_line(line: str) -> str:
    # entfernt z.B. '01. [exact]' am Zeilenanfang
    line = re.sub(r"^\s*\d+\.\s*\[[^\]]+\]\s*", "", line)
    return line.strip().strip('"').strip()


def _is_noise_line(line: str) -> bool:
    if not line:
        return True
    if _LABEL_RE.search(line):
        return True
    if _RELDATE_INLINE_RE.search(line):
        return True
    return False


def _extract_from_multiline_title(raw_title: str):
    """
    Robustere Heuristik für jobs.json title:
    - title enthält oft Sammeltext: Zeit, Jobtitel, Labels, Ort, Firma.
    - Wir filtern Labels/relative Zeiten auch wenn inline.
    - Jobtitel = erste non-noise Zeile.
    - Firma = letzte non-noise Zeile mit Rechtsform (AG/GmbH/SA/...) sonst letzte non-noise Zeile.
    - Ort = Zeile nach "Arbeitsort:" falls vorhanden, sonst erste non-noise Zeile mit City-Hint.
    """
    raw_lines = [_normalize_line(x) for x in (raw_title or "").splitlines()]
    raw_lines = [x for x in raw_lines if x]

    # location: explizit nach "Arbeitsort"
    location = ""
    for i, line in enumerate(raw_lines):
        if line.lower().startswith("arbeitsort"):
            if i + 1 < len(raw_lines):
                location = _normalize_line(raw_lines[i + 1])
            break

    clean = [line for line in raw_lines if not _is_noise_line(line)]

    job_title = clean[0] if clean else ""
    company = ""

    # Firma: letzte Zeile mit Rechtsform-Hint
    for line in reversed(clean):
        if _COMPANY_HINT_RE.search(line):
            company = line
            break

    # fallback: letzte clean Zeile (wenn nicht schon job_title)
    if not company and len(clean) >= 2:
        company = clean[-1]
        if company == job_title:
            company = ""

    # fallback location via city hint
    if not location:
        for line in clean[1:]:
            if _CITY_HINT_RE.search(line):
                location = line
                break

    if location == company:
        location = ""

    return job_title, company, location


def _select_template(title: str, templates_dir: Path) -> Path:
    t = (title or "").lower()
    if any(
        k in t
        for k in [
            "logistik",
            "lager",
            "kommission",
            "versand",
            "wareneingang",
            "warenausgang",
        ]
    ):
        p = templates_dir / "T3_Logistik.docx"
        if p.exists():
            return p
    if any(
        k in t
        for k in ["system", "techniker", "engineer", "operator", "netzw", "noc"]
    ):
        p = templates_dir / "T2_Systemtechnik.docx"
        if p.exists():
            return p
    return templates_dir / "T1_ITSupport.docx"


def _replace_placeholders_docx(doc, mapping: dict):
    # paragraphs
    for p in doc.paragraphs:
        for run in p.runs:
            txt = run.text
            for k, v in mapping.items():
                if k in txt:
                    txt = txt.replace(k, v)
            run.text = txt
    # tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        txt = run.text
                        for k, v in mapping.items():
                            if k in txt:
                                txt = txt.replace(k, v)
                        run.text = txt


def _send_mail_with_attachments(to_addr, subject, body_text, attachments):
    """
    Versendet eine E-Mail mit Anhängen via SMTP.
    attachments: Liste von Pfaden (Path-Objekte oder Strings).
    """
    from config import config

    msg = MIMEMultipart()
    msg["From"] = config.SENDER_EMAIL
    msg["To"] = to_addr
    msg["Subject"] = subject
    if os.getenv("SMTP_BCC"):
        msg["Bcc"] = os.getenv("SMTP_BCC")

    msg.attach(MIMEText(body_text, "plain", "utf-8"))

    for fpath in attachments:
        fpath = Path(fpath)
        if not fpath.exists():
            print(f"WARNUNG: Anhang fehlt: {fpath}")
            continue
        with open(fpath, "rb") as f:
            part = MIMEApplication(f.read(), Name=fpath.name)
        part.add_header("Content-Disposition", "attachment", filename=fpath.name)
        msg.attach(part)

    # SMTP Send
    try:
        server = smtplib.SMTP(config.SMTP_SERVER, int(config.SMTP_PORT))
        server.starttls()
        server.login(config.SENDER_EMAIL, config.SENDER_PASSWORD)
        server.send_message(msg)
        server.quit()
        return True
    except Exception as e:
        print(f"SMTP FEHLER bei {to_addr}: {e}")
        return False


def _find_application_doc(out_dir: Path, company: str, job_title: str) -> Path:
    """
    Versucht, das passende DOCX in out/ zu finden.
    Strategie: Suche nach Dateien, die company und job_title (sanitized) enthalten.
    """
    safe_comp = _sanitize_filename(company)
    safe_title = _sanitize_filename(job_title)

    # 1. Versuch: Exakter Match des Standard-Musters
    candidates = list(out_dir.glob(f"*{safe_comp}*{safe_title}*.docx"))
    if not candidates:
        # 2. Versuch: Nur Company (falls Titel abweicht)
        candidates = list(out_dir.glob(f"*{safe_comp}*.docx"))

    if not candidates:
        return None

    # Nimm das neueste
    candidates.sort(key=lambda p: p.stat().st_mtime, reverse=True)
    return candidates[0]


def cmd_send_applications(args):
    """
    Versendet vorbereitete Bewerbungen per E-Mail.
    Benötigt:
    - data/jobs.json
    - out/ (generierte DOCX)
    - Bewerbungsunterlagen/ (CV, Zeugnisse)
    """
    # 1. Checks & Config
    if os.getenv("SEND_APPLICATIONS_ENABLED", "false").lower() not in {
        "true",
        "1",
        "yes",
    }:
        print("Versand deaktiviert. Setze SEND_APPLICATIONS_ENABLED=true in .env")
        return

    limit = int(os.getenv("DAILY_SEND_LIMIT", "5"))
    proj = Path(args.proj).resolve() if args.proj else Path.cwd()
    in_path = Path(args.in_file) if args.in_file else (proj / "data" / "jobs.json")
    out_dir = Path(args.out_dir) if args.out_dir else (proj / "out")
    tracker_path = (
        Path(args.tracker)
        if args.tracker
        else (proj / "bewerbungen_tracking.csv")
    )

    # Anhänge (statisch)
    cv_path = proj / "Bewerbungsunterlagen" / "Florian Bujupi Lebenslauf.pdf"
    certs_path = proj / "Bewerbungsunterlagen" / "Zeugnisse kom. - Florian Bujupi.pdf"

    if not cv_path.exists():
        print(f"ABBRUCH: Lebenslauf fehlt: {cv_path}")
        return

    # 2. Load Jobs
    if not in_path.exists():
        print("Keine jobs.json gefunden.")
        return
    jobs = json.loads(in_path.read_text(encoding="utf-8"))

    # 3. Iterate
    sent_count = 0
    print(f"Starte Versand (Limit: {limit})...")

    for job in jobs:
        if sent_count >= limit:
            print("Tageslimit erreicht.")
            break

        if job.get("fit") != "OK":
            continue

        to_addr = job.get("application_email")
        if not to_addr:
            continue

        company = job.get("company", "Firma")
        title = job.get("title", "Job")

        # Doc finden
        docx_path = _find_application_doc(out_dir, company, title)
        if not docx_path:
            print(f"Skip {company}: Kein Anschreiben in {out_dir} gefunden.")
            continue

        # Email bauen
        subject = f"Bewerbung als {title} - Florian Bujupi"

        salutation = "Sehr geehrte Damen und Herren"

        body = f"""{salutation}

anbei erhalten Sie meine Bewerbung für die Position als {title}.

Besonders an {company} reizt mich die ausgeschriebene Position und die Möglichkeit, meine Erfahrungen im IT-Support und der Systemadministration gewinnbringend einzubringen.

Im Anhang finden Sie mein Anschreiben, den Lebenslauf sowie meine Zeugnisse.

Für ein persönliches Gespräch stehe ich Ihnen gerne zur Verfügung.

Freundliche Grüsse
Florian Bujupi
"""

        attachments = [docx_path, cv_path]
        if certs_path.exists():
            attachments.append(certs_path)

        # Senden
        print(f"Sende an {to_addr} ({company})...")
        if args.dry_run:
            print("  [DRY RUN] Mail wäre gesendet worden.")
            success = True
        else:
            success = _send_mail_with_attachments(
                to_addr, subject, body, attachments
            )

        if success:
            sent_count += 1
            # Archive
            cmd_archive_sent(
                argparse.Namespace(file=str(docx_path), company=company, dest="")
            )
            # Tracker Update
            with open(tracker_path, "a", encoding="utf-8") as f:
                today = datetime.now().strftime("%d.%m.%Y")
                f.write(
                    f'{today},"{company}","{title}","EMAIL","{to_addr}","VERSENDET",""\n'
                )

    print(f"Versand abgeschlossen. {sent_count} E-Mails gesendet.")


def cmd_prepare_applications(args):
    """
    Liest data/jobs.json, nimmt fit=="OK" (oder --force-all),
    wendet .docx Templates an, schreibt out/*.docx und
    hängt neue Zeilen an bewerbungen_tracking.csv.
    """
    from job_collector import compute_fit
    auto_fit = str(os.getenv("AUTO_FIT_ENABLED", "false")).lower() in {
        "1",
        "true",
        "t",
        "yes",
        "y",
        "ja",
        "j",
    }
    min_score_apply = float(os.getenv("MIN_SCORE_APPLY", "1") or 1)
    from docx import Document

    proj = Path(args.proj).resolve() if args.proj else Path.cwd()
    in_path = Path(args.in_file) if args.in_file else (proj / "data" / "jobs.json")
    out_dir = Path(args.out_dir) if args.out_dir else (proj / "out")
    templates_dir = (
        Path(args.templates_dir)
        if args.templates_dir
        else (proj / "Anschreiben_Templates")
    )
    tracker_path = (
        Path(args.tracker)
        if args.tracker
        else (proj / "bewerbungen_tracking.csv")
    )

    if not in_path.exists():
        print(f"FEHLER: {in_path} nicht gefunden.")
        raise SystemExit(1)
    if not templates_dir.exists():
        print(f"FEHLER: Templates-Ordner fehlt: {templates_dir}")
        raise SystemExit(1)
    if not out_dir.exists():
        print(f"FEHLER: out/ fehlt: {out_dir} (Ordner bitte einmal anlegen).")
        raise SystemExit(1)

    jobs = json.loads(in_path.read_text(encoding="utf-8"))

    header = "Datum,Firma,Position,Portal,Link,Status,Notizen\n"
    if not tracker_path.exists():
        tracker_path.write_text(header, encoding="utf-8")
    else:
        txt = tracker_path.read_text(encoding="utf-8")
        if not txt.strip().startswith("Datum,"):
            tracker_path.write_text(header + txt, encoding="utf-8")

    today = datetime.now().strftime("%d.%m.%Y")
    stamp = datetime.now().strftime("%Y%m%d")

    prepared = 0
    sent_base = Path(args.copy_sent_dir) if args.copy_sent_dir else None
    if args.mirror_sent and not sent_base:
        sent_base = proj / "04_Versendete_Bewerbungen"
    if sent_base and not sent_base.exists():
        sent_base.mkdir(parents=True, exist_ok=True)
    for job in jobs:
        fit = (job.get("fit") or "").upper()
        if auto_fit:
            score_val = job.get("score") or 0
            try:
                score_val = float(score_val)
            except Exception:
                score_val = 0
            fit = compute_fit(job.get("match", ""), score_val, min_score_apply)
            job["fit"] = fit
        if not args.force_all and fit != "OK":
            continue

        raw_title = job.get("title", "")
        job_title = job.get("job_title") or job.get("position") or ""
        company = job.get("company") or ""
        location = job.get("location") or ""

        if (not job_title) or (not company):
            t2, c2, l2 = _extract_from_multiline_title(raw_title)
            if not job_title and t2:
                job_title = t2
            if not company and c2:
                company = c2
            if not location and l2:
                location = l2

        if not company:
            company = "Firma Unbekannt"
        if not job_title:
            job_title = "Position Unbekannt"

        source = job.get("source") or job.get("portal") or ""
        url = job.get("url") or job.get("link") or ""

        template_path = _select_template(job_title, templates_dir)
        if not template_path.exists():
            print(f"FEHLER: Template fehlt: {template_path}")
        continue

        doc = Document(str(template_path))

        # Mapping für neue Tokens plus bisherige Platzhalter
        mapping = {
            "{{TODAY_DATE}}": today,
            "{{JOB_TITLE}}": job_title,
            "{{COMPANY_NAME}}": company,
            "{{SALUTATION}}": "Sehr geehrte Damen und Herren",
            "{{COMPANY_HOOK_1SENT}}": f"mich Ihr Unternehmen {company} und die ausgeschriebene Position sehr ansprechen.",
            "{{AD_MATCH_2TO3_SENTENCES}}": (
                "Besonders meine Erfahrung im 1st/2nd Level Support und "
                "mein Verständnis für logistische Prozesse kann ich gewinnbringend einbringen."
            ),
            "<Ort>": location or "Bülach",
            "<Datum>": today,
            "<JOBTITEL>": job_title,
            "<FIRMA>": company,
        }
        # Support für alte Tokens mit EINZELNEN Klammern
        legacy_mapping = {
            "{TODAY_DATE}": today,
            "{JOB_TITLE}": job_title,
            "{COMPANY_NAME}": company,
            "{SALUTATION}": "Sehr geehrte Damen und Herren",
        }
        mapping.update(legacy_mapping)

        # Falls schon konkrete Bezeichner wie "{Evergreen Human Resources AG}" im DOCX stehen:
        mapping[f"{{{company}}}"] = company

        _replace_placeholders_docx(doc, mapping)

        out_name = _sanitize_filename(f"{company}_{job_title}_{stamp}.docx")
        out_path = out_dir / out_name
        doc.save(str(out_path))

        # Optional: Kopie in 04_Versendete_Bewerbungen/<Firma>/...
        if sent_base:
            target_dir = sent_base / _sanitize_filename(company or "Unbekannt")
            target_dir.mkdir(parents=True, exist_ok=True)
            shutil.copy2(out_path, target_dir / out_name)

        row = (
            f'{today},"{company}","{job_title}","{source}","{url}","Erstellt",""\n'
        )
        with tracker_path.open("a", encoding="utf-8") as f:
            f.write(row)

        print(f"Erstellt: {out_name}")
        prepared += 1

    print(f"Fertig. {prepared} Bewerbungen vorbereitet.")


def main(argv=None):
    p = argparse.ArgumentParser()
    sub = p.add_subparsers(dest="cmd", required=True)

    sub.add_parser("env-check")
    sub.add_parser("verify")
    arch = sub.add_parser("archive-sent")
    arch.add_argument(
        "--file", required=True, help="Pfad zur versendeten DOCX (z.B. aus out/)"
    )
    arch.add_argument(
        "--company", default="", help="Optional Firmenname override"
    )
    arch.add_argument(
        "--dest",
        default="",
        help="Basis-Ordner fuer Kopie (default: 04_Versendete_Bewerbungen)",
    )
    sub.add_parser("gen-templates")
    sub.add_parser("start")
    sub.add_parser("open")
    sub.add_parser("email-test")
    sub.add_parser("list")
    mail = sub.add_parser("mail-list")
    mail.add_argument(
        "--dry-run", action="store_true", help="Nur simulieren, keine Mails senden"
    )

    sub.add_parser("tracker-sync")

    mark_applied = sub.add_parser("mark-applied")
    mark_applied.add_argument("job_uid", nargs="?", help="Job UID")
    mark_applied.add_argument("--url", default="", help="Job URL")

    mark_ignored = sub.add_parser("mark-ignored")
    mark_ignored.add_argument("job_uid", nargs="?", help="Job UID")
    mark_ignored.add_argument("--url", default="", help="Job URL")

    prep = sub.add_parser("prepare-applications")
    prep.add_argument("--proj", default="", help="Projekt-Root (default: cwd)")
    prep.add_argument("--in", dest="in_file", default="", help="Input jobs.json")
    prep.add_argument(
        "--out", dest="out_dir", default="", help="Output-Ordner out/"
    )
    prep.add_argument(
        "--templates",
        dest="templates_dir",
        default="",
        help="Templates-Ordner",
    )
    prep.add_argument("--tracker", default="", help="Tracker CSV")
    prep.add_argument(
        "--force-all", action="store_true", help="Alle Jobs verarbeiten, egal fit"
    )
    prep.add_argument(
        "--mirror-sent",
        action="store_true",
        help=(
            "Optional Kopie der erzeugten Anschreiben in "
            "04_Versendete_Bewerbungen/<Firma>/ ablegen"
        ),
    )
    prep.add_argument(
        "--copy-sent-dir",
        default="",
        help=(
            "Alternativer Basis-Ordner fuer die Kopien (default: "
            "04_Versendete_Bewerbungen im Projekt)"
        ),
    )

    send = sub.add_parser("send-applications")
    send.add_argument("--proj", default="", help="Projekt-Root")
    send.add_argument("--in", dest="in_file", default="", help="Input jobs.json")
    send.add_argument(
        "--out", dest="out_dir", default="", help="Output-Ordner out/"
    )
    send.add_argument("--tracker", default="", help="Tracker CSV")
    send.add_argument(
        "--dry-run", action="store_true", help="Nur simulieren, keine Mails senden"
    )

    args = p.parse_args(argv)

    if args.cmd == "env-check":
        cmd_env_check(args)
    elif args.cmd == "gen-templates":
        cmd_gen_templates(args)
    elif args.cmd == "start":
        cmd_start(args)
    elif args.cmd == "open":
        cmd_open(args)
    elif args.cmd == "email-test":
        cmd_email_test(args)
    elif args.cmd == "list":
        cmd_list(args)
    elif args.cmd == "mail-list":
        cmd_mail_list(args)
    elif args.cmd == "tracker-sync":
        cmd_tracker_sync(args)
    elif args.cmd == "mark-applied":
        cmd_mark_applied(args)
    elif args.cmd == "mark-ignored":
        cmd_mark_ignored(args)
    elif args.cmd == "prepare-applications":
        cmd_prepare_applications(args)
    elif args.cmd == "send-applications":
        cmd_send_applications(args)


if __name__ == "__main__":
    main()
