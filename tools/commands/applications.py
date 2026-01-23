from __future__ import annotations

import json
import os
import re
import shutil
from datetime import datetime
from pathlib import Path
from typing import Iterable

import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication

from bewerbungsagent.job_text_utils import extract_from_multiline_title
from tools.common import is_dry_run


def _sanitize_filename(value: str) -> str:
    cleaned = (value or "").strip()
    cleaned = re.sub(r"[\\/:*?\"<>|]", "_", cleaned)
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned[:120] if len(cleaned) > 120 else cleaned


def _select_template(title: str, templates_dir: Path) -> Path:
    t = (title or "").lower()
    if any(k in t for k in ["logistik", "lager", "kommission", "versand", "wareneingang", "warenausgang"]):
        candidate = templates_dir / "T3_Logistik.docx"
        if candidate.exists():
            return candidate
    if any(k in t for k in ["system", "techniker", "engineer", "operator", "netzw", "noc"]):
        candidate = templates_dir / "T2_Systemtechnik.docx"
        if candidate.exists():
            return candidate
    return templates_dir / "T1_ITSupport.docx"


def _replace_placeholders_docx(doc, mapping: dict) -> None:
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            text = run.text
            for key, value in mapping.items():
                if key in text:
                    text = text.replace(key, value)
            run.text = text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        text = run.text
                        for key, value in mapping.items():
                            if key in text:
                                text = text.replace(key, value)
                        run.text = text


def _find_application_doc(out_dir: Path, company: str, job_title: str) -> Path | None:
    safe_comp = _sanitize_filename(company)
    safe_title = _sanitize_filename(job_title)
    candidates = list(out_dir.glob(f"*{safe_comp}*{safe_title}*.docx"))
    if not candidates:
        candidates = list(out_dir.glob(f"*{safe_comp}*.docx"))
    if not candidates:
        return None
    candidates.sort(key=lambda p: p.stat().st_mtime, reverse=True)
    return candidates[0]


def _relpath_if_possible(path: Path) -> str:
    try:
        return str(path.resolve().relative_to(Path.cwd().resolve()))
    except Exception:
        return str(path)


def _send_mail_with_attachments(to_addr: str, subject: str, body_text: str, attachments: Iterable[Path]) -> bool:
    from bewerbungsagent.config import config

    msg = MIMEMultipart()
    msg["From"] = config.SENDER_EMAIL
    msg["To"] = to_addr
    msg["Subject"] = subject
    if os.getenv("SMTP_BCC"):
        msg["Bcc"] = os.getenv("SMTP_BCC")

    msg.attach(MIMEText(body_text, "plain", "utf-8"))

    for fpath in attachments:
        fpath = Path(fpath)
        if not fpath.exists():
            print(f"WARNUNG: Anhang fehlt: {fpath}")
            continue
        with open(fpath, "rb") as fh:
            part = MIMEApplication(fh.read(), Name=fpath.name)
        part.add_header("Content-Disposition", "attachment", filename=fpath.name)
        msg.attach(part)

    try:
        server = smtplib.SMTP(config.SMTP_SERVER, int(config.SMTP_PORT))
        server.starttls()
        server.login(config.SENDER_EMAIL, config.SENDER_PASSWORD)
        server.send_message(msg)
        server.quit()
        return True
    except Exception as exc:
        print(f"SMTP FEHLER bei {to_addr}: {exc}")
        return False


def _prepare_tracker_header(tracker_path: Path) -> None:
    header = "Datum,Firma,Position,Portal,Link,Status,Notizen\n"
    if not tracker_path.exists():
        tracker_path.write_text(header, encoding="utf-8")
        return
    existing = tracker_path.read_text(encoding="utf-8")
    if not existing.strip().startswith("Datum,"):
        tracker_path.write_text(header + existing, encoding="utf-8")


def _job_fit(job: dict, auto_fit: bool, min_score_apply: float) -> str:
    if not auto_fit:
        return (job.get("fit") or "").upper()
    from bewerbungsagent.job_collector import compute_fit

    score_val = job.get("score") or 0
    try:
        score_val = float(score_val)
    except Exception:
        score_val = 0
    fit = compute_fit(job.get("match", ""), score_val, min_score_apply)
    job["fit"] = fit
    return fit


def _resolve_job_fields(job: dict) -> tuple[str, str, str, str, str]:
    raw_title = job.get("title", "")
    job_title = job.get("job_title") or job.get("position") or ""
    company = job.get("company") or ""
    location = job.get("location") or ""
    if (not job_title) or (not company):
        t2, c2, l2 = extract_from_multiline_title(raw_title)
        if not job_title and t2:
            job_title = t2
        if not company and c2:
            company = c2
        if not location and l2:
            location = l2
    if not company:
        company = "Firma Unbekannt"
    if not job_title:
        job_title = "Position Unbekannt"
    source = job.get("source") or job.get("portal") or ""
    url = job.get("url") or job.get("link") or ""
    return job_title, company, location, source, url


def _build_mapping(today: str, job_title: str, company: str, location: str) -> dict:
    mapping = {
        "{{TODAY_DATE}}": today,
        "{{JOB_TITLE}}": job_title,
        "{{COMPANY_NAME}}": company,
        "{{SALUTATION}}": "Sehr geehrte Damen und Herren",
        "{{COMPANY_HOOK_1SENT}}": f"mich Ihr Unternehmen {company} und die ausgeschriebene Position sehr ansprechen.",
        "{{AD_MATCH_2TO3_SENTENCES}}": (
            "Besonders meine Erfahrung im 1st/2nd Level Support und "
            "mein Verstaendnis fuer logistische Prozesse kann ich gewinnbringend einbringen."
        ),
        "<Ort>": location or "Buelach",
        "<Datum>": today,
        "<JOBTITEL>": job_title,
        "<FIRMA>": company,
    }
    legacy = {
        "{TODAY_DATE}": today,
        "{JOB_TITLE}": job_title,
        "{COMPANY_NAME}": company,
        "{SALUTATION}": "Sehr geehrte Damen und Herren",
    }
    mapping.update(legacy)
    mapping[f"{{{company}}}"] = company
    return mapping


def prepare_applications(args) -> None:
    from docx import Document

    auto_fit = str(os.getenv("AUTO_FIT_ENABLED", "false")).lower() in {
        "1",
        "true",
        "t",
        "yes",
        "y",
        "ja",
        "j",
    }
    min_score_apply = float(os.getenv("MIN_SCORE_APPLY", "1") or 1)

    proj = Path(args.proj).resolve() if args.proj else Path.cwd()
    in_path = Path(args.in_file) if args.in_file else (proj / "data" / "jobs.json")
    out_dir = Path(args.out_dir) if args.out_dir else (proj / "out")
    templates_dir = (
        Path(args.templates_dir)
        if args.templates_dir
        else (proj / "Anschreiben_Templates")
    )
    tracker_path = (
        Path(args.tracker)
        if args.tracker
        else (proj / "data" / "bewerbungen_tracking.csv")
    )

    if not in_path.exists():
        print(f"FEHLER: {in_path} nicht gefunden.")
        raise SystemExit(1)
    if not templates_dir.exists():
        print(f"FEHLER: Templates-Ordner fehlt: {templates_dir}")
        raise SystemExit(1)
    if not out_dir.exists():
        print(f"FEHLER: out/ fehlt: {out_dir} (Ordner bitte einmal anlegen).")
        raise SystemExit(1)

    jobs = json.loads(in_path.read_text(encoding="utf-8"))
    _prepare_tracker_header(tracker_path)

    today = datetime.now().strftime("%d.%m.%Y")
    stamp = datetime.now().strftime("%Y%m%d")

    prepared = 0
    sent_base = Path(args.copy_sent_dir) if args.copy_sent_dir else None
    if args.mirror_sent and not sent_base:
        sent_base = proj / "04_Versendete_Bewerbungen"
    if sent_base and not sent_base.exists():
        sent_base.mkdir(parents=True, exist_ok=True)

    for job in jobs:
        fit = _job_fit(job, auto_fit, min_score_apply)
        if not args.force_all and fit != "OK":
            continue

        job_title, company, location, source, url = _resolve_job_fields(job)
        template_path = _select_template(job_title, templates_dir)
        if not template_path.exists():
            print(f"FEHLER: Template fehlt: {template_path}")
            continue

        doc = Document(str(template_path))
        mapping = _build_mapping(today, job_title, company, location)
        _replace_placeholders_docx(doc, mapping)

        out_name = _sanitize_filename(f"{company}_{job_title}_{stamp}.docx")
        out_path = out_dir / out_name
        doc.save(str(out_path))

        if sent_base:
            target_dir = sent_base / _sanitize_filename(company or "Unbekannt")
            target_dir.mkdir(parents=True, exist_ok=True)
            shutil.copy2(out_path, target_dir / out_name)

        row = f'{today},"{company}","{job_title}","{source}","{url}","Erstellt",""\n'
        with tracker_path.open("a", encoding="utf-8") as fh:
            fh.write(row)

        print(f"Erstellt: {out_name}")
        prepared += 1

    print(f"Fertig. {prepared} Bewerbungen vorbereitet.")


def _archive_sent_file(src: Path, company: str, dest_base: Path | None) -> Path:
    base = dest_base if dest_base else (Path.cwd() / "04_Versendete_Bewerbungen")
    target_dir = base / _sanitize_filename(company)
    target_dir.mkdir(parents=True, exist_ok=True)
    dest = target_dir / src.name
    shutil.copy2(src, dest)
    return dest


def archive_sent(args) -> Path:
    src = Path(args.file).expanduser()
    if not src.exists():
        print(f"FEHLER: Datei nicht gefunden: {src}")
        raise SystemExit(1)

    company = args.company
    if not company:
        parts = src.stem.split("_")
        if parts:
            company = parts[0]
    company = company or "Unbekannt"

    dest_base = Path(args.dest).expanduser() if args.dest else None
    dest = _archive_sent_file(src, company, dest_base)
    print(f"Kopie erstellt: {dest}")
    return dest


def _update_state_after_send(job: dict, docx_path: Path, archive_dest: Path | None) -> None:
    try:
        from bewerbungsagent.job_state import (
            STATUS_APPLIED,
            build_job_uid,
            canonicalize_url,
            load_state,
            now_iso,
            save_state,
        )

        stamp = now_iso()
        state = load_state()
        job_uid, canonical = build_job_uid(job)
        record = state.get(job_uid, {})
        link = job.get("url") or job.get("link") or ""
        record.update(
            {
                "job_uid": job_uid,
                "source": job.get("source") or record.get("source") or "",
                "canonical_url": canonical
                or record.get("canonical_url")
                or canonicalize_url(link)
                or link,
                "link": link or record.get("link") or "",
                "title": job.get("title") or record.get("title") or "",
                "company": job.get("company") or record.get("company") or "",
                "location": job.get("location") or record.get("location") or "",
                "first_seen_at": record.get("first_seen_at") or stamp,
                "last_seen_at": stamp,
                "last_sent_at": stamp,
                "status": STATUS_APPLIED,
                "application_doc": _relpath_if_possible(docx_path),
                "application_doc_archived": _relpath_if_possible(archive_dest)
                if archive_dest
                else "",
                "application_sent_at": stamp,
            }
        )
        state[job_uid] = record
        save_state(state)
        try:
            from bewerbungsagent.job_tracker import (
                get_tracker_path,
                load_tracker,
                write_tracker,
            )

            tracker_rows = load_tracker(get_tracker_path())
            write_tracker(state, get_tracker_path(), tracker_rows)
        except Exception:
            pass
    except Exception:
        pass


def send_applications(args) -> None:
    if str(os.getenv("SEND_APPLICATIONS_ENABLED", "false")).lower() not in {
        "true",
        "1",
        "yes",
    }:
        print("Versand deaktiviert. Setze SEND_APPLICATIONS_ENABLED=true in .env")
        return

    limit = int(os.getenv("DAILY_SEND_LIMIT", "5"))
    proj = Path(args.proj).resolve() if args.proj else Path.cwd()
    in_path = Path(args.in_file) if args.in_file else (proj / "data" / "jobs.json")
    out_dir = Path(args.out_dir) if args.out_dir else (proj / "out")
    tracker_path = (
        Path(args.tracker)
        if args.tracker
        else (proj / "data" / "bewerbungen_tracking.csv")
    )

    cv_path = proj / "Bewerbungsunterlagen" / "Florian Bujupi Lebenslauf.pdf"
    certs_path = proj / "Bewerbungsunterlagen" / "Zeugnisse kom. - Florian Bujupi.pdf"
    if not cv_path.exists():
        print(f"ABBRUCH: Lebenslauf fehlt: {cv_path}")
        return

    if not in_path.exists():
        print("Keine jobs.json gefunden.")
        return
    jobs = json.loads(in_path.read_text(encoding="utf-8"))

    sent_count = 0
    print(f"Starte Versand (Limit: {limit})...")

    for job in jobs:
        if sent_count >= limit:
            print("Tageslimit erreicht.")
            break
        if job.get("fit") != "OK":
            continue

        to_addr = job.get("application_email")
        if not to_addr:
            continue

        company = job.get("company", "Firma")
        title = job.get("title", "Job")

        docx_path = _find_application_doc(out_dir, company, title)
        if not docx_path:
            print(f"Skip {company}: Kein Anschreiben in {out_dir} gefunden.")
            continue

        subject = f"Bewerbung als {title} - Florian Bujupi"
        salutation = "Sehr geehrte Damen und Herren"
        body = f"""{salutation}

anbei erhalten Sie meine Bewerbung fuer die Position als {title}.

Besonders an {company} reizt mich die ausgeschriebene Position und die Moeglichkeit, meine Erfahrungen im IT-Support und der Systemadministration gewinnbringend einzubringen.

Im Anhang finden Sie mein Anschreiben, den Lebenslauf sowie meine Zeugnisse.

Fuer ein persoenliches Gespraech stehe ich Ihnen gerne zur Verfuegung.

Freundliche Gruesse
Florian Bujupi
"""

        attachments = [docx_path, cv_path]
        if certs_path.exists():
            attachments.append(certs_path)

        print(f"Sende an {to_addr} ({company})...")
        if is_dry_run(args):
            print("  [DRY RUN] Mail waere gesendet worden.")
            sent_count += 1
            continue

        success = _send_mail_with_attachments(to_addr, subject, body, attachments)
        if not success:
            continue

        sent_count += 1
        archive_dest = _archive_sent_file(docx_path, company, None)
        _update_state_after_send(job, docx_path, archive_dest)

        with open(tracker_path, "a", encoding="utf-8") as fh:
            today = datetime.now().strftime("%d.%m.%Y")
            fh.write(
                f'{today},"{company}","{title}","EMAIL","{to_addr}","VERSENDET",""\n'
            )

    print(f"Versand abgeschlossen. {sent_count} E-Mails gesendet.")
