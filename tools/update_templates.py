import os
from pathlib import Path
from docx import Document

def create_template(path, title_key, hook_token, match_token):
    doc = Document()
    
    # Header
    doc.add_paragraph("Florian Bujupi")
    doc.add_paragraph("Bülach, Schweiz")
    doc.add_paragraph("fb95.jobs@gmail.com")
    doc.add_paragraph("")
    doc.add_paragraph("Winterthur, {{TODAY_DATE}}")
    doc.add_paragraph("")
    
    # Subject
    p = doc.add_paragraph()
    runner = p.add_run(f"Bewerbung als {title_key} bei {{COMPANY_NAME}}")
    runner.bold = True
    
    doc.add_paragraph("")
    
    # Body
    doc.add_paragraph("{{SALUTATION}}")
    doc.add_paragraph("")
    doc.add_paragraph(f"ich bewerbe mich als {title_key} bei {{COMPANY_NAME}}, weil {hook_token}.")
    doc.add_paragraph("")
    
    # Main Text (Standard block)
    text = (
        "In meinen bisherigen Rollen habe ich zuverlässig in strukturierten IT- und Prozessumgebungen gearbeitet "
        "(u. a. SAP-nahe Abläufe, Benutzer-/Rollenlogik, Daten- und Prozessqualität). Ich bringe solide Grundlagen "
        "in Netzwerken, Systemadministration, Windows-Client-Support sowie Python/SQL mit und bin es gewohnt, "
        "Tickets sauber zu analysieren, Prioritäten zu setzen und Lösungen nachvollziehbar zu dokumentieren. "
        f"{match_token}"
    )
    doc.add_paragraph(text)
    doc.add_paragraph("")
    
    doc.add_paragraph(
        "Ich arbeite ruhig, exakt und serviceorientiert, lerne schnell neue Tools/Stacks und übernehme Verantwortung, "
        "ohne lange Einarbeitung zu benötigen. Start ist sofort möglich; Pensum 80–100 % bevorzugt, ab 60 % möglich."
    )
    doc.add_paragraph("")
    doc.add_paragraph("Über eine Einladung zum Gespräch freue ich mich.")
    doc.add_paragraph("")
    doc.add_paragraph("Freundliche Grüsse")
    doc.add_paragraph("Florian Bujupi")
    
    doc.save(path)
    print(f"Created {path}")

def main():
    base = Path("Anschreiben_Templates")
    base.mkdir(exist_ok=True)
    
    # T1 IT Support
    create_template(
        base / "T1_ITSupport.docx", 
        "{{JOB_TITLE}}", 
        "{{COMPANY_HOOK_1SENT}}", 
        "{{AD_MATCH_2TO3_SENTENCES}}"
    )
    
    # T2 Systemtechnik (Variant)
    create_template(
        base / "T2_Systemtechnik.docx", 
        "{{JOB_TITLE}}", 
        "{{COMPANY_HOOK_1SENT}}", 
        "Besonders meine Erfahrung in der Systemadministration und Netzwerktechnik möchte ich hier einbringen."
    )

    # T3 Logistik (Variant)
    create_template(
        base / "T3_Logistik.docx", 
        "{{JOB_TITLE}}", 
        "{{COMPANY_HOOK_1SENT}}", 
        "Meine Kombination aus IT-Verständnis und Logistik-Erfahrung (SAP, Prozesse) passt ideal zu dieser Stelle."
    )

if __name__ == "__main__":
    main()
